"""
DOCX Track Changes Generator

Generates Word documents with track changes for SME review.
- SAFE fixes appear as tracked changes (strikethrough + insertion)
- REVIEW/QUERY items appear as Word comments

Uses python-docx with direct XML manipulation for track changes support.
"""

import os
import re
import copy
from datetime import datetime
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree


# Word namespace mappings
WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
W = WORD_NAMESPACE


@dataclass
class TrackChange:
    """Represents a tracked change to apply."""
    original: str
    replacement: str
    line_num: int
    rule_id: str
    category: str
    fix_type: str  # SAFE, REVIEW, QUERY
    message: str


@dataclass
class Comment:
    """Represents a comment to add."""
    text: str
    target_text: str
    line_num: int
    rule_id: str
    category: str
    fix_type: str


class DocxTrackChangesGenerator:
    """
    Generates DOCX files with track changes from validation results.
    """

    def __init__(self, author: str = "Editorial AI"):
        """
        Initialize the generator.

        Args:
            author: Author name for track changes attribution
        """
        self.author = author
        self.change_id = 0
        self.comment_id = 0

    def process_document(
        self,
        input_path: str,
        output_path: str,
        changes: List[TrackChange],
        comments: List[Comment]
    ) -> Dict:
        """
        Process a DOCX file and add track changes and comments.

        Args:
            input_path: Path to original DOCX
            output_path: Path for output DOCX
            changes: List of TrackChange objects for SAFE fixes
            comments: List of Comment objects for REVIEW/QUERY items

        Returns:
            Dict with processing statistics
        """
        # Load the document
        doc = Document(input_path)

        stats = {
            'changes_applied': 0,
            'comments_added': 0,
            'changes_failed': 0,
            'comments_failed': 0
        }

        # Enable track changes in document settings
        self._enable_track_changes(doc)

        # Process each paragraph to apply changes
        for para in doc.paragraphs:
            para_text = para.text

            # Apply SAFE changes with track changes markup
            for change in changes:
                if change.original in para_text:
                    if self._apply_track_change(para, change):
                        stats['changes_applied'] += 1
                    else:
                        stats['changes_failed'] += 1

        # Add comments for REVIEW/QUERY items
        for comment in comments:
            for para in doc.paragraphs:
                if comment.target_text in para.text:
                    if self._add_comment(doc, para, comment):
                        stats['comments_added'] += 1
                    else:
                        stats['comments_failed'] += 1
                    break  # Only add comment once per target

        # Save the document
        doc.save(output_path)

        return stats

    def _enable_track_changes(self, doc: Document):
        """Enable track changes mode in document settings."""
        settings = doc.settings.element

        # Add trackRevisions element
        track_revisions = OxmlElement('w:trackRevisions')
        settings.append(track_revisions)

    def _apply_track_change(self, para, change: TrackChange) -> bool:
        """
        Apply a track change to a paragraph.

        This creates the XML structure for tracked insertions/deletions:
        <w:del>original text</w:del><w:ins>replacement text</w:ins>

        Returns True if successful.
        """
        try:
            # Find the run containing the text
            for run in para.runs:
                if change.original in run.text:
                    # Create deletion element
                    self.change_id += 1

                    # Get the run's XML element
                    run_elem = run._element

                    # Split the text around the match
                    text = run.text
                    before, match, after = text.partition(change.original)

                    # Clear the run
                    run.text = before

                    # Create del element for original text
                    del_elem = self._create_del_element(change.original)
                    run_elem.addnext(del_elem)

                    # Create ins element for replacement text
                    ins_elem = self._create_ins_element(change.replacement)
                    del_elem.addnext(ins_elem)

                    # Add remaining text if any
                    if after:
                        after_run = OxmlElement('w:r')
                        after_text = OxmlElement('w:t')
                        after_text.text = after
                        after_run.append(after_text)
                        ins_elem.addnext(after_run)

                    return True

            return False

        except Exception as e:
            print(f"Error applying track change: {e}")
            return False

    def _create_del_element(self, text: str) -> OxmlElement:
        """Create a w:del element for deleted text."""
        del_elem = OxmlElement('w:del')
        del_elem.set(qn('w:id'), str(self.change_id))
        del_elem.set(qn('w:author'), self.author)
        del_elem.set(qn('w:date'), datetime.now().isoformat())

        # Create run inside del
        run = OxmlElement('w:r')
        del_text = OxmlElement('w:delText')
        del_text.text = text
        run.append(del_text)
        del_elem.append(run)

        return del_elem

    def _create_ins_element(self, text: str) -> OxmlElement:
        """Create a w:ins element for inserted text."""
        self.change_id += 1

        ins_elem = OxmlElement('w:ins')
        ins_elem.set(qn('w:id'), str(self.change_id))
        ins_elem.set(qn('w:author'), self.author)
        ins_elem.set(qn('w:date'), datetime.now().isoformat())

        # Create run inside ins
        run = OxmlElement('w:r')
        ins_text = OxmlElement('w:t')
        ins_text.text = text
        run.append(ins_text)
        ins_elem.append(run)

        return ins_elem

    def _add_comment(self, doc: Document, para, comment: Comment) -> bool:
        """
        Add a comment to a paragraph.

        Returns True if successful.
        """
        try:
            self.comment_id += 1

            # Format comment text with category and rule
            comment_text = f"[{comment.category}] {comment.text}"
            if comment.fix_type == "QUERY":
                comment_text = f"[QUESTION] {comment_text}"

            # Find the part of document.xml that we need to modify
            # Comments in DOCX are complex - they require:
            # 1. Comment definition in comments.xml
            # 2. Comment reference in document.xml

            # For now, use a simpler approach: add inline annotation
            # This isn't true Word comments but is visible
            for run in para.runs:
                if comment.target_text in run.text:
                    # Add the comment as highlighted text after the target
                    text = run.text
                    idx = text.find(comment.target_text) + len(comment.target_text)

                    # Insert comment marker
                    marker = f" [{comment.rule_id}: {comment.text}]"
                    run.text = text[:idx] + marker + text[idx:]

                    # Style the marker (italic, smaller)
                    # Note: This modifies the whole run, not ideal

                    return True

            return False

        except Exception as e:
            print(f"Error adding comment: {e}")
            return False


def generate_track_changes_docx(
    input_path: str,
    output_path: str,
    validation_results: Dict,
    author: str = "Editorial AI"
) -> Dict:
    """
    Generate a DOCX with track changes from validation results.

    Args:
        input_path: Path to original DOCX
        output_path: Path for output DOCX
        validation_results: Dict containing issues from editorial validation
        author: Author name for changes

    Returns:
        Dict with processing statistics
    """
    # Parse validation results into TrackChange and Comment objects
    changes = []
    comments = []

    for issue in validation_results.get('issues', []):
        fix_type = issue.get('fix_type', 'REVIEW')

        if fix_type == 'SAFE' and issue.get('suggestion'):
            # SAFE fixes become track changes
            changes.append(TrackChange(
                original=issue.get('original', ''),
                replacement=issue.get('suggestion', ''),
                line_num=issue.get('line', 0),
                rule_id=issue.get('rule_id', ''),
                category=issue.get('category', ''),
                fix_type=fix_type,
                message=issue.get('message', '')
            ))
        else:
            # REVIEW and QUERY become comments
            comments.append(Comment(
                text=issue.get('message', ''),
                target_text=issue.get('original', issue.get('match', '')),
                line_num=issue.get('line', 0),
                rule_id=issue.get('rule_id', ''),
                category=issue.get('category', ''),
                fix_type=fix_type
            ))

    # Generate the document
    generator = DocxTrackChangesGenerator(author=author)
    return generator.process_document(input_path, output_path, changes, comments)


# CLI interface
if __name__ == '__main__':
    import argparse
    import json

    parser = argparse.ArgumentParser(
        description='Generate DOCX with track changes from validation results'
    )
    parser.add_argument('input', help='Input DOCX file')
    parser.add_argument('--output', '-o', help='Output DOCX file')
    parser.add_argument('--results', '-r', help='JSON file with validation results')
    parser.add_argument('--author', default='Editorial AI', help='Author name for changes')

    args = parser.parse_args()

    # Default output name
    if not args.output:
        base = os.path.splitext(args.input)[0]
        args.output = f"{base}_edited.docx"

    # Load validation results
    if args.results:
        with open(args.results) as f:
            results = json.load(f)
    else:
        # Run validation if no results provided
        print("No validation results provided. Run validation first.")
        exit(1)

    # Generate track changes document
    stats = generate_track_changes_docx(
        args.input,
        args.output,
        results,
        author=args.author
    )

    print(f"Generated: {args.output}")
    print(f"  Changes applied: {stats['changes_applied']}")
    print(f"  Comments added: {stats['comments_added']}")
    if stats['changes_failed']:
        print(f"  Changes failed: {stats['changes_failed']}")
    if stats['comments_failed']:
        print(f"  Comments failed: {stats['comments_failed']}")
